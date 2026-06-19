# backend/tasks/geotech_tasks.py
"""
Celery task: generate monthly environmental compliance report (DOCX).

Checks effluent values against IFC PS3 regulatory limits:
  CN WAD   ≤ 50 mg/L
  Total CN ≤ 100 mg/L
  As       ≤ 0.5 mg/L
  Hg       ≤ 0.01 mg/L

Report saved as DOCX in uploads/reports/{project_id}/env_report_{month}.docx
"""
from __future__ import annotations
import os
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    from celery_app import celery_app
    from db import conn, release
except ImportError:
    from backend.celery_app import celery_app
    from backend.db import conn, release

IFC_PS3_LIMITS = {
    "cn_wad_mg_l":   50.0,
    "total_cn_mg_l": 100.0,
    "as_mg_l":       0.5,
    "hg_mg_l":       0.01,
}

PARAM_LABELS = {
    "cn_wad_mg_l":   "CN WAD (mg/L)",
    "total_cn_mg_l": "Total CN (mg/L)",
    "as_mg_l":       "As (mg/L)",
    "hg_mg_l":       "Hg (mg/L)",
}


@celery_app.task(name="tasks.geotech_tasks.generate_env_report")
def generate_env_report(project_id: str, month: str) -> str:
    """
    Generate monthly IFC PS3 compliance report.

    Args:
        project_id: UUID of the project
        month: YYYY-MM string (e.g. '2026-04')

    Returns:
        Path to the generated DOCX file.
    """
    from docx import Document

    db = None
    try:
        db = conn()
        with db.cursor() as cur:
            cur.execute(
                """SELECT tag_name, AVG(value) as avg_value,
                          MAX(value) as max_value, COUNT(*) as sample_count
                   FROM tag_readings tr
                   JOIN process_tags pt ON pt.id = tr.tag_id
                   WHERE pt.project_id = %s
                     AND to_char(tr.time, 'YYYY-MM') = %s
                     AND pt.tag_name = ANY(%s)
                   GROUP BY tag_name""",
                (project_id, month, list(IFC_PS3_LIMITS.keys()))
            )
            rows = cur.fetchall()
    except Exception:
        if db is not None:
            db.rollback()
        logger.exception("Failed to query tag readings for env report project=%s month=%s", project_id, month)
        raise
    finally:
        if db is not None:
            release(db)

    doc = Document()
    doc.add_heading("Monthly Environmental Compliance Report", 0)
    doc.add_paragraph(f"Project ID: {project_id}")
    doc.add_paragraph(f"Reporting Period: {month}")
    doc.add_paragraph("Standard: IFC Performance Standard 3 (PS3)")
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()} UTC")

    doc.add_heading("Effluent Quality — IFC PS3 Limits", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Parameter", "Limit", "Monthly Avg", "Monthly Max", "Status"]):
        hdr[i].text = h

    results_by_param = {r[0]: r for r in rows}
    for param, limit in IFC_PS3_LIMITS.items():
        row = table.add_row().cells
        row[0].text = PARAM_LABELS.get(param, param)
        row[1].text = str(limit)
        if param in results_by_param:
            _, avg, mx, _ = results_by_param[param]
            row[2].text = f"{avg:.4f}"
            row[3].text = f"{mx:.4f}"
            status = "COMPLIANT" if mx <= limit else "EXCEEDANCE"
            row[4].text = status
        else:
            row[2].text = "No data"
            row[3].text = "No data"
            row[4].text = "—"

    out_dir = os.path.join("uploads", "reports", project_id)
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"env_report_{month}.docx")
    doc.save(out_path)
    logger.info("Generated env report: %s", out_path)
    return out_path
