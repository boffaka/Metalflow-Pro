"""
NI 43-101 Report Export — PDF (weasyprint) and Word (python-docx).
"""
from __future__ import annotations
import html as html_mod
import io
import logging
from datetime import datetime

logger = logging.getLogger("mpdpms.ni43101_export")


def _esc(text: str) -> str:
    return html_mod.escape(text or "")


# ─── Word Export (python-docx) ────────────────────────────────────────────────

def generate_docx(sections: list[dict], lang: str, project: dict) -> bytes:
    """Generate a .docx file from NI 43-101 sections. Returns raw bytes."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Title page
    title_field = "title_fr" if lang == "fr" else "title_en"
    content_field = "content_fr" if lang == "fr" else "content_en"

    proj_name = project.get("project_name", "N/D") if project else "N/D"
    proj_code = project.get("project_code", "") if project else ""

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NI 43-101 Technical Report" if lang == "en" else "Rapport Technique NI 43-101")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0xD4, 0xA0, 0x1A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(proj_name)
    run.font.size = Pt(18)

    if proj_code:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Code: {proj_code}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%Y-%m-%d"))
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # Table of contents placeholder
    toc_title = "Table des matieres" if lang == "fr" else "Table of Contents"
    doc.add_heading(toc_title, level=1)
    for s in sections:
        key = s.get("subsection_key", s.get("key", ""))
        title = s.get(title_field, "")
        p = doc.add_paragraph(f"{key}  {title}", style="List Number")
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # Sections
    current_section = None
    for s in sections:
        sec_num = s.get("section_number", 13)
        key = s.get("subsection_key", s.get("key", ""))
        title = s.get(title_field, "")
        content = s.get(content_field, "")

        # Section header
        if sec_num != current_section:
            current_section = sec_num
            if sec_num == 13:
                heading = ("Section 13 — Essais metallurgiques et traitement mineral"
                           if lang == "fr" else
                           "Section 13 — Mineral Processing and Metallurgical Testing")
            else:
                heading = ("Section 17 — Methodes de recuperation"
                           if lang == "fr" else
                           "Section 17 — Recovery Methods")
            doc.add_heading(heading, level=1)

        # Subsection heading
        doc.add_heading(f"{key} — {title}", level=2)

        # Content — process line by line for basic markdown-like formatting
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("- "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip("*"))
                run.bold = True
            else:
                doc.add_paragraph(line)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── PDF Export (xhtml2pdf) ───────────────────────────────────────────────────

_PDF_CSS = """
@page {
    size: A4;
    margin: 2.5cm 2.5cm 3cm 3cm;
}
body {
    font-family: Helvetica, Arial, sans-serif;
    font-size: 11pt; line-height: 1.5; color: #222;
}
h1 { font-size: 18pt; color: #1a365d; border-bottom: 2px solid #d4a01a;
     padding-bottom: 6pt; margin-top: 24pt; }
h2 { font-size: 13pt; color: #2d3748; margin-top: 14pt; }
.title-page { text-align: center; padding-top: 100pt; }
.title-page h1 { border: none; color: #d4a01a; font-size: 24pt; }
.title-page .project { font-size: 16pt; color: #222; margin-top: 12pt; }
.title-page .date { font-size: 11pt; color: #888; margin-top: 24pt; }
ul { margin-left: 18pt; }
li { margin-bottom: 2pt; }
.page-break { page-break-before: always; }
"""


def generate_pdf(sections: list[dict], lang: str, project: dict) -> bytes:
    """Generate a PDF from NI 43-101 sections using xhtml2pdf. Returns raw bytes."""
    from xhtml2pdf import pisa

    title_field = "title_fr" if lang == "fr" else "title_en"
    content_field = "content_fr" if lang == "fr" else "content_en"

    proj_name = _esc(project.get("project_name", "N/D") if project else "N/D")
    proj_code = _esc(project.get("project_code", "") if project else "")

    html_parts = [
        f"<!DOCTYPE html><html><head><meta charset='utf-8'><style>{_PDF_CSS}</style></head><body>",
        f'<div class="title-page">',
        f'<h1>{"Rapport Technique NI 43-101" if lang == "fr" else "NI 43-101 Technical Report"}</h1>',
        f'<div class="project">{proj_name}</div>',
    ]
    if proj_code:
        html_parts.append(f'<div style="color:#888;margin-top:8pt">Code: {proj_code}</div>')
    html_parts.append(f'<div class="date">{datetime.now().strftime("%Y-%m-%d")}</div>')
    html_parts.append('</div>')

    current_section = None
    for s in sections:
        sec_num = s.get("section_number", 13)
        key = _esc(s.get("subsection_key", s.get("key", "")))
        title = _esc(s.get(title_field, ""))
        content = s.get(content_field, "")

        if sec_num != current_section:
            current_section = sec_num
            if sec_num == 13:
                heading = ("Section 13 — Essais metallurgiques et traitement mineral"
                           if lang == "fr" else
                           "Section 13 — Mineral Processing and Metallurgical Testing")
            else:
                heading = ("Section 17 — Methodes de recuperation"
                           if lang == "fr" else
                           "Section 17 — Recovery Methods")
            html_parts.append(f'<div class="page-break"></div>')
            html_parts.append(f'<h1>{_esc(heading)}</h1>')

        html_parts.append(f'<h2>{key} — {title}</h2>')

        # Convert content to HTML
        lines = content.split("\n")
        in_list = False
        for line in lines:
            line = line.strip()
            if not line:
                if in_list:
                    html_parts.append("</ul>")
                    in_list = False
                continue
            if line.startswith("- "):
                if not in_list:
                    html_parts.append("<ul>")
                    in_list = True
                html_parts.append(f"<li>{_esc(line[2:])}</li>")
            else:
                if in_list:
                    html_parts.append("</ul>")
                    in_list = False
                if line.startswith("**") and line.endswith("**"):
                    html_parts.append(f"<p><strong>{_esc(line.strip('*'))}</strong></p>")
                else:
                    html_parts.append(f"<p>{_esc(line)}</p>")
        if in_list:
            html_parts.append("</ul>")

    html_parts.append("</body></html>")
    html_str = "\n".join(html_parts)

    buf = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html_str), dest=buf)
    buf.seek(0)
    return buf.getvalue()
